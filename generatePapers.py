from docx import Document
from docx.shared import Pt
from enum import Enum, unique
import copy
import csv
import re


class Student(object):
    """
        Student class
        Property:name,schoolId,sujects
    """

    def __init__(self, name, school_id):
        self.__name = name
        self.__schoolId = school_id

    def __str__(self):
        string = "姓名：%s\n学号：%s\n" % (self.__name, self.__schoolId)
        string += str(self.__subjects)
        return string

    def __repr__(self):
        string = "姓名：%s\n学号：%s\n" % (self.__name, self.__schoolId)
        string += str(self.__subjects)
        return string

    @property
    def name(self):
        return self.__name

    @name.setter
    def name(self, name):
        self.__name = name

    @property
    def schoolId(self):
        return self.__schoolId

    @schoolId.setter
    def schoolId(self, school_id):
        self.__schoolId = school_id

    @property
    def subjects(self):
        return self.__subjects

    @subjects.setter
    def subjects(self, subjects):
        self.__subjects = copy.deepcopy(subjects)


@unique
class SubjectType(Enum):
    SelectSubject = 0
    JudgmentSubject = 1
    ProgramSubject = 2
    Completion = 3


class Subject(object):

    def __init__(self, Type, title, options, answer, right_answer):
        self.__type = Type
        self.__title = title
        self.__options = options
        self.__answer = answer
        self.__right_answer = right_answer

    def __str__(self):
        string = ""

        # 题目类型
        if self.__type == SubjectType.SelectSubject:
            string += "题目类型：选择题\n"
        elif self.__type == SubjectType.ProgramSubject:
            string += "题目类型：程序题\n"
        elif self.__type == SubjectType.Completion:
            string += "题目类型：填空题\n"
        elif self.__type == SubjectType.JudgmentSubject:
            string += "题目类型：判断题\n"

        # 标题
        string += "题目：%s\n" % self.__title

        # 选项
        string += "选项：%s\n" % str(self.__options)

        # 正确答案
        string += "正确答案：%s\n" % self.__right_answer

        # 考生答案
        string += "考生答案：%s\n" % self.__answer
        return string

    def __repr__(self):
        return self.__str__()

    @property
    def type(self):
        return self.__type

    @type.setter
    def type(self, mtype):
        self.__type = copy(mtype)

    @property
    def title(self):
        return self.__title

    @title.setter
    def title(self, title):
        self.__title = copy(title)

    @property
    def options(self):
        return self.__options

    @options.setter
    def options(self, options):
        self.__options = copy.deepcopy(options)

    @property
    def answer(self):
        return self.__answer

    @answer.setter
    def answer(self, answer):
        self.__answer = answer

    @property
    def right_answer(self):
        return self.__right_answer

    @right_answer.setter
    def right_answer(self, right_answer):
        self.__right_answer = right_answer


# set the student's information
def setStudentInfo(filename):
    global students
    csvfile = open(filename, 'r', encoding='utf8')
    csv.reader('\xEF\xBB\xBF')
    reader = csv.reader(csvfile)
    last_school_id = ""
    subjects = []
    student = None
    for item in reader:

        # 忽略第一行
        if reader.line_num == 1:
            continue

        # 对题目进行分类
        subject_type = None
        if item[6] != '' and item[13] in 'ABCD':
            subject_type = SubjectType.SelectSubject
            # print("选择题")
            # print(item[2])
        elif item[12] == 'A' or item[12] == 'B':
            subject_type = SubjectType.JudgmentSubject
            # print("判断题")
            # print(item[2])
        elif '填空' in item[2]:
            subject_type = SubjectType.Completion
            # print("填空题")
            # print(item[2])
        else:
            subject_type = SubjectType.ProgramSubject
            # print("程序题")
            # print(item[2])
        if item[1] != last_school_id:
            if last_school_id != "":
                student.subjects = copy.deepcopy(subjects)
                students.append(copy.deepcopy(student))
                subjects.clear()
            student = Student(item[0], item[1])
            last_school_id = item[1]
        if subject_type == SubjectType.SelectSubject:
            subjects.append(Subject(subject_type, replace_html_tag(item[2]),
                                    [replace_html_tag(item[3]), replace_html_tag(item[4]), replace_html_tag(item[5]), \
                                     replace_html_tag(item[6]), ], item[13], item[12]))
        else:
            subjects.append(Subject(subject_type, replace_html_tag(item[2]), [], replace_html_tag(item[13]),replace_html_tag(item[12])))
            
    student.subjects = copy.deepcopy(subjects)
    students.append(copy.deepcopy(student))
    csvfile.close()


def generatePapers(document, student):
    doc = document
    tables = doc.tables
    tables[1].cell(0, 1).text = student.name
    tables[1].cell(1, 1).text = student.schoolId
    subjects = student.subjects
    select_subject_count = True
    program_subject_count = True
    judgment_subject_count = True
    completion_count = True
    count = 1
    for subject in subjects:
        # 大题名
        if select_subject_count and subject.type == SubjectType.SelectSubject:
            doc.add_paragraph().add_run('第1类:单选题')
            select_subject_count = False

        if judgment_subject_count and subject.type == SubjectType.JudgmentSubject:
            doc.add_paragraph().add_run('第2类: 是非题')
            judgment_subject_count = False

        if completion_count and subject.type == SubjectType.Completion:
            doc.add_paragraph().add_run('第3类: 填空题')
            completion_count = False

        if program_subject_count and subject.type == SubjectType.ProgramSubject:
            doc.add_paragraph().add_run('第4类: 问答题')
            completion_count = False

        # 处理单选题
        if subject.type == SubjectType.SelectSubject:
            table = doc.add_table(rows=6, cols=1)
            r = table.cell(0, 0).add_paragraph().add_run('%s、%s' % (count, subject.title))
            r.font.name = 'Trebuchet MS'
            r.bold = True
            r.size = Pt(10.5)
            options = subject.options
            table.cell(0, 1).text = 'A.%s' % options[0]
            table.cell(0, 2).text = 'B.%s' % options[1]
            table.cell(0, 3).text = 'C.%s' % options[2]
            table.cell(0, 4).text = 'D.%s' % options[3]
            right = '√' if subject.right_answer == subject.answer else '×'
            fraction = 2.0 if subject.right_answer == subject.answer else 0.0
            table.cell(0, 5).text = '标准答案:%s   考生答案:%s   %s   得分%s' % (subject.right_answer, \
                                                                       subject.answer, right, fraction)
        # 处理是非题
        elif subject.type == SubjectType.JudgmentSubject:
            table = doc.add_table(rows=2, cols=1)
            r = table.cell(0, 0).add_paragraph().add_run('%s、%s' % (count, subject.title))
            r.font.name = 'Trebuchet MS'
            r.bold = True
            r.size = Pt(10.5)
            right = '√' if subject.right_answer == subject.answer else '×'
            fraction = 2.0 if subject.right_answer == subject.answer else 0.0
            table.cell(0, 1).text = '标准答案:%s   考生答案:%s   %s   得分%s' % (subject.right_answer, \
                                                                       subject.answer, right, fraction)
        # 处理填空题
        elif subject.type == SubjectType.Completion:
            answers = opCompleteAnswer(subject.answer)
            print(answers)
            rights = []
            right_answer = opCompleteAnswer(subject.right_answer)
            print(right_answer)
            perfraction = 4/3.0
            fraction = 0
            for i in range(len(right_answer)):
                right = '√' if i >= len(answers) or right_answer[i] == answers[i] else '×'
                rights.append(right)
                fraction = fraction + perfraction
            table = doc.add_table(rows = 3, cols = 1)
            r = table.cell(0, 0).add_paragraph().add_run('%s、%s' % (count, subject.title))
            r.font.name = 'Trebuchet MS'
            r.bold = True
            r.size = Pt(10.5)
            text = '标准答案: '
            for string in right_answer:
                text = text + ' ' + string
            text = text + '  考生答案:'
            for string in answers:
                text = text + ' ' + string
            table.cell(0,1).text = text
            text = ''
            for i in range(len(right_answer)):
                text = text + '第%s空：%s ' % (i+1,rights[i])
            text = text + ' 得分 %s' % fraction
            table.cell(0,2).text = text
            
        else:
            table = doc.add_table(rows = 3, cols = 1)
            r = table.cell(0, 0).add_paragraph().add_run('%s、%s' % (count, subject.title))
            r.font.name = 'Trebuchet MS'
            r.bold = True
            r.size = Pt(10.5)
            table.cell(0,1).text = '标准答案:%s' % subject.right_answer
            table.cell(0,2).text = '考生答案:%s' % subject.answer
            
        count += 1
    doc.save("%s%s.doc" % (student.schoolId,student.name))

# 去除html各种标签
def replace_html_tag(string):
    # replace方式替换标签
    string = string.replace('<br>', '\n')
    string = string.replace('<br/>', '\n')
    string = string.replace('<br >', '\n')
    string = string.replace('<br />', '\n')
    string = string.replace('&quot;', '"')
    string = string.replace('&gt;', '>')
    string = string.replace('&lt;', '<')
    string = string.replace('&amp;', '&')
    string = string.replace('&nbsp;', ' ')
    string = string.strip()
    return string

def opCompleteAnswer(string):
    string = string.strip()
    return string.split('@`_~@')

# 所有学生的list
students = []
document = None
if __name__ == '__main__':
    # open template file
    document = Document('template.docx')
    # document.save('C:\\Users\\Administrator.WIN-N52B1L3VP08\\Desktop\\test.docx')
    setStudentInfo('siti.csv')
    for student in students:
        generatePapers(copy.deepcopy(document),student)